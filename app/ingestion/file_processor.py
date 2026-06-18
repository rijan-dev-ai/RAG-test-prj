"""
File upload ingestion: extract text from PDF/DOCX/TXT, chunk, embed, store.
"""
import os
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional


from app.core.logging import get_logger
from app.ingestion.chunking import chunk_text
from app.ingestion.web_scraper import _store_chunks
from app.vectorstore.base import BaseVectorStore

logger = get_logger(__name__)

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


def extract_text_from_pdf(file_path: str) -> str:
    from pypdf import PdfReader

    reader = PdfReader(file_path)
    pages_text = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages_text.append(f"[Page {i + 1}]\n{page_text}")
    return "\n\n".join(pages_text)


def extract_text_from_docx(file_path: str) -> str:
    import docx

    document = docx.Document(file_path)
    parts = []
    for para in document.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also pull text from tables
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


def extract_text_from_txt(file_path: str) -> str:
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(file_path: str) -> str:
    """Dispatch text extraction based on file extension."""
    ext = Path(file_path).suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    elif ext in (".txt", ".md"):
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(
            f"Unsupported file type: {ext}. Supported types: {sorted(SUPPORTED_EXTENSIONS)}"
        )


def ingest_file(
    file_path: str,
    vector_store: BaseVectorStore,
    original_filename: Optional[str] = None,
    skip_duplicates: bool = True,
    extra_metadata: Optional[dict] = None,
) -> dict:
    """Full pipeline: extract -> chunk -> embed -> store for an uploaded file.

    Returns a summary dict with counts of chunks added/skipped.
    """
    filename = original_filename or os.path.basename(file_path)
    logger.info("Ingesting file: %s", filename)

    text = extract_text(file_path)
    if not text.strip():
        logger.warning("No text extracted from %s", filename)
        return {"filename": filename, "chunks_added": 0, "chunks_skipped": 0}

    metadata = {
        "source": filename,
        "source_type": "file",
        "file_extension": Path(filename).suffix.lower(),
        "ingested_at": datetime.now(timezone.utc).isoformat(),
    }
    if extra_metadata:
        metadata.update(extra_metadata)

    chunks = chunk_text(text, metadata)
    added, skipped = _store_chunks(chunks, vector_store, skip_duplicates)

    logger.info("File %s: %d chunks added, %d skipped (duplicates)", filename, added, skipped)
    return {
        "filename": filename,
        "chunks_added": added,
        "chunks_skipped": skipped,
        "total_chunks": len(chunks),
    }
